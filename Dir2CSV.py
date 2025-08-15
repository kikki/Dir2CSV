# Dir2CSV.py
import os
import csv
from pathlib import Path
from typing import List, Dict, Optional
import threading
from typing import Iterable, Sequence

TO_Scan_Ordner = ""  # (kept for compatibility; not used directly)
CSV_AUSGABE_DATEI = "scanresult.csv"

# === Optional imports for PDF/DOCX text extraction ===
try:
    import PyPDF2  # Fast and robust for many PDFs
except Exception:
    PyPDF2 = None

try:
    # pdfminer.six – more accurate with complex layouts
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except Exception:
    pdfminer_extract_text = None

try:
    import docx  # python-docx
except Exception:
    docx = None


# -----------------------------------------------------------------------------
# Class: FileScanner
# Purpose: Scan a directory for files, gather basic metadata and readable text
#          content (for supported types), and export results to a CSV file.
# -----------------------------------------------------------------------------
class FileScanner:
    """
    Scans files in a directory and stores the results in a CSV file.
    """

    # -------------------------------------------------------------------------
    # __init__
    # Purpose: Initialize file-type sets and exclusion rules.
    # -------------------------------------------------------------------------
    def __init__(self):
        # File extensions we are interested in
        self.target_extensions = {
            '.cs', '.xml', '.csproj', '.sln', '.dll',
            '.json', '.py', '.md', '.jt', '.jtx', '.js', '.jsx',
            '.txt', '.log', '.cfg', '.conf', '.ini', '.properties',
            '.yaml', '.yml', '.html', '.htm', '.csv', '.tsv', '.rst',
            '.asc', '.ascx', '.xaml', '.css', '.scss', '.less', '.sql',
            '.php', '.rb', '.java',
            '.pdf', '.docx'
        }

        # File extensions we will try to open and read as text
        self.readable_extensions = {
            '.cs', '.xml', '.csproj', '.json', '.py', '.md', '.jt', '.jtx', '.js', '.jsx',
            '.txt', '.log', '.cfg', '.conf', '.ini', '.properties', '.yaml', '.yml',
            '.html', '.htm', '.csv', '.tsv', '.rst', '.asc', '.ascx', '.xaml',
            '.css', '.scss', '.less', '.sql', '.php', '.rb', '.java',
            '.pdf', '.docx'
        }

        # Binary file types that we do NOT attempt to read as text
        self.binary_extensions = {
            '.dll', '.exe', '.dxt', '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico',
            '.zip', '.rar', '.tar', '.gz', '.7z', '.bin', '.dat', '.lib', '.pdb',
            '.class', '.jar', '.war', '.ear', '.apk', '.ipa'
            # IMPORTANT: do not add .pdf/.docx here because we extract their text
        }

        # Folder names to skip when reading content
        self.excluded_folder_names_for_content = {'.git', 'node_modules', 'venv','.venv', '__pycache__','dist', 'build', 'out', 'target', 'bin', 'obj','Dir2CSV_config'}

        # Optional: content length limit per file (disabled in this version)
        # self.max_chars_per_file = max_chars_per_file

    # -------------------------------------------------------------------------
    # scan_folder
    # Purpose: Walk through the given folder recursively and collect file info.
    # Returns: List of dictionaries with file metadata and (optional) content.
    # -------------------------------------------------------------------------
    def scan_folder(self, start_folder: str) -> List[Dict]:
        found_files = []
        self.start_path = Path(start_folder).resolve()  # absolute path for relative computation

        print(f"Scanning folder: {self.start_path}")

        for file_path in self.start_path.rglob('*'):
            if file_path.is_file():
                ext = file_path.suffix.lower()
                if ext in self.target_extensions or self.should_check_file(file_path):
                    info = self.collect_file_info(file_path)
                    if info:
                        found_files.append(info)
                    else:
                        print(f"******ERROR******* File could not be processed: {file_path}")

        print(f"Total files found: {len(found_files)}")
        return found_files

    # -------------------------------------------------------------------------
    # should_check_file
    # Purpose: Decide whether a file without a known extension should be probed.
    # Logic: Skip huge files and dot/tilde-prefixed names; otherwise check.
    # -------------------------------------------------------------------------
    def should_check_file(self, file_path: Path) -> bool:
        ext = file_path.suffix.lower()
        if not ext or ext not in self.target_extensions:
            try:
                if file_path.stat().st_size > 10 * 1024 * 1024:  # >10MB => skip
                    return False
            except:
                return False

            if file_path.name.startswith('.') or file_path.name.startswith('~'):
                return False

            return True
        return False

    # -------------------------------------------------------------------------
    # collect_file_info
    # Purpose: Build a dictionary with relative path, name, extension, and
    #          (if applicable) extracted text content.
    # -------------------------------------------------------------------------
    def collect_file_info(self, file_path: Path) -> Optional[Dict]:
        try:
            rel_to_start = file_path.relative_to(self.start_path)
            relative_path = Path(self.start_path.name) / rel_to_start

            ext = file_path.suffix.lower()
            name = file_path.name

            info = {
                'relative_path': str(relative_path),
                'file_name': name,
                'file_extension': ext,
                'content': ''
            }

            if not self.is_in_excluded_folder(file_path):
                if ext in self.binary_extensions:
                    info['content'] = '[Binary file – content not readable]'
                elif ext in self.readable_extensions:
                    if ext == '.pdf':
                        text = self.read_pdf(file_path)
                    elif ext == '.docx':
                        text = self.read_docx(file_path)
                    else:
                        text = self.read_file_content(file_path)
                    info['content'] = self._csv_safe(text)
                else:
                    text = self.read_file_intelligently(file_path)
                    info['content'] = self._csv_safe(text)

            print(f"Processed: {name}")
            return info

        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return None

    # -------------------------------------------------------------------------
    # is_in_excluded_folder
    # Purpose: Check whether a file is located in a content-excluded directory.
    # -------------------------------------------------------------------------
    def is_in_excluded_folder(self, file_path: Path) -> bool:
        #for _triple in self.excluded_folder_names_for_content:
        #    if str(file_path).lower().index(_triple.lower()) >= 0:
        #        return True
        #return False
        for part in file_path.parts:
            if part.lower() in self.excluded_folder_names_for_content:
                print(f"++is_in_excluded_folder : {file_path} return TRUE -> Bleibt Leer")
                return True
        print(f"++is_in_excluded_folder : {file_path} return False -> wird auf Text geprüft")
        return False

    # -------------------------------------------------------------------------
    # read_file_intelligently
    # Purpose: Heuristically detect if a file is text; if so, read as text,
    #          otherwise mark as binary-like content.
    # -------------------------------------------------------------------------
    def read_file_intelligently(self, file_path: Path) -> str:
        try:
            with open(file_path, 'rb') as f:
                first_bytes = f.read(1024)
            if self.is_probably_text(first_bytes):
                return self.read_file_content(file_path)
            else:
                return '[Binary-like file detected – content not readable]'
        except Exception as e:
            return f"[Error during intelligent read: {e}]"

    # -------------------------------------------------------------------------
    # is_probably_text
    # Purpose: Simple heuristic to check if a byte buffer looks like text.
    # -------------------------------------------------------------------------
    def is_probably_text(self, byte_data: bytes) -> bool:
        if not byte_data:
            return False
        if b'\x00' in byte_data:
            return False
        printable = 0
        for b in byte_data:
            if 32 <= b <= 126 or b in [9, 10, 13]:
                printable += 1
        ratio = printable / len(byte_data)
        return ratio >= 0.8

    # -------------------------------------------------------------------------
    # read_file_content
    # Purpose: Try multiple encodings to read a text file and return its content.
    # -------------------------------------------------------------------------
    def read_file_content(self, file_path: Path) -> str:
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc, errors='strict') as f:
                    text = f.read()
                    return self._truncate(text)
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return f"[Error reading file: {e}]"
        return "[Error: Could not read file due to encoding issues]"

    # -------------------------------------------------------------------------
    # read_pdf
    # Purpose: Extract text from a PDF using PyPDF2 (fast) or pdfminer.six (fallback).
    # -------------------------------------------------------------------------
    def read_pdf(self, file_path: Path) -> str:
        """
        Extract text from a PDF file. Tries PyPDF2 first, then pdfminer.six.
        """
        # 1) PyPDF2 (fast)
        if PyPDF2 is not None:
            try:
                parts = []
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        t = page.extract_text() or ''
                        parts.append(t)
                text = "\n".join(parts).strip()
                if text:
                    return self._truncate(text)
            except Exception:
                pass  # fall back to pdfminer

        # 2) pdfminer.six (more accurate)
        if pdfminer_extract_text is not None:
            try:
                text = pdfminer_extract_text(str(file_path)) or ''
                return self._truncate(text)
            except Exception as e:
                return f"[PDF read error (pdfminer): {e}]"

        return "[Note: No PDF text library available. Install 'PyPDF2' or 'pdfminer.six']"

    # -------------------------------------------------------------------------
    # read_docx
    # Purpose: Extract text from a DOCX file using python-docx.
    # -------------------------------------------------------------------------
    def read_docx(self, file_path: Path) -> str:
        """
        Extract text from a DOCX file via python-docx.
        """
        if docx is None:
            return "[Note: 'python-docx' not installed. Please run 'pip install python-docx'.]"
        try:
            d = docx.Document(str(file_path))
            parts = []
            # Paragraphs
            for p in d.paragraphs:
                parts.append(p.text)
            # Table cells (often important)
            for t in d.tables:
                for row in t.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            text = "\n".join(parts).strip()
            return self._truncate(text)
        except Exception as e:
            return f"[DOCX read error: {e}]"

    # -------------------------------------------------------------------------
    # _truncate
    # Purpose: Optionally shorten extracted text to a maximum length.
    #          (Disabled here but kept for future use.)
    # -------------------------------------------------------------------------
    def _truncate(self, text: str) -> str:
        if text is None:
            return ""
        # if len(text) > self.max_chars_per_file:
        #     return text[: self.max_chars_per_file] + f"\n[...truncated to {self.max_chars_per_file} chars...]"
        return text

    # -------------------------------------------------------------------------
    # _csv_safe
    # Purpose: Make content safe for CSV (escape quotes, visualize newlines/tabs).
    # -------------------------------------------------------------------------
    def _csv_safe(self, content: str) -> str:
        if content is None:
            content = ""
        # Escape double quotes
        content = content.replace('"', '""')
        # Normalize newlines and tabs
        content = content.replace('\r\n', '\\n').replace('\n', '\\n').replace('\r', '\\n')
        content = content.replace('\t', '\\t')
        return content

    # -------------------------------------------------------------------------
    # create_csv
    # Purpose: Write the collected file info into a semicolon-separated CSV file.
    # -------------------------------------------------------------------------
    def create_csv(self, file_list: List[Dict], output_file: str):
        if not file_list:
            print("No files to save.")
            return

        try:
            with open(output_file, 'w', newline='', encoding='utf-8-sig') as csv_file:
                columns = ['relative_path', 'file_name', 'file_extension', 'content']
                writer = csv.DictWriter(
                    csv_file,
                    fieldnames=columns,
                    delimiter=',',
                    quotechar='"',
                    quoting=csv.QUOTE_ALL
                )
                writer.writeheader()
                for info in file_list:
                    writer.writerow(info)

            print(f"CSV successfully created: {output_file}")
            print(f"Entry count: {len(file_list)}")

        except Exception as e:
            print(f"Error creating CSV: {e}")
if __name__ == "__main__":
    scanner =FileScanner()
    found = scanner.scan_folder(TO_Scan_Ordner)
    scanner.create_csv(found, CSV_AUSGABE_DATEI)
    print( f"Scan finished.\nFile saved to:\n{CSV_AUSGABE_DATEI}")
    